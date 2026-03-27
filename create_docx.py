import sys
import subprocess
import os

def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])
    finally:
        globals()[package] = __import__(package)

install_and_import('docx')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Add Title
title = doc.add_heading('TÀI LIỆU HƯỚNG DẪN SỬ DỤNG', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)

def add_borders(table):
    tbl = table._tbl # get xml element in table
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    
    tblPr = tbl.tblPr
    tblPr.append(tblBorders)

# 1. Trang web dành cho chủ đầu tư
p1 = doc.add_paragraph()
r1 = p1.add_run('1. Trang web dành cho chủ đầu tư')
r1.bold = True
r1.font.size = Pt(12)

table1 = doc.add_table(rows=1, cols=2)
hdr_cells1 = table1.rows[0].cells
hdr_cells1[0].text = 'Chức năng/Bước thực hiện'
hdr_cells1[1].text = 'Màn hình demo'
for cell in hdr_cells1:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

data1 = [
    ("Đăng nhập vào hệ thống\n- Bước 1: Nhập địa chỉ hệ thống trên trình duyệt.\n- Bước 2: Nhập Tài khoản & Mật khẩu.\n- Bước 3: Nhấn nút [Đăng nhập].", "[Chèn ảnh màn hình đăng nhập tại đây]"),
    ("Quản lý dự án\n- Bước 1: Truy cập menu Quản lý dự án.\n- Bước 2: Click [Thêm mới] để tạo dự án hoặc chọn một dự án có sẵn để xem và chỉnh sửa.\n- Bước 3: Điền các thông tin của dự án, sau đó click [Lưu].", "[Chèn ảnh màn hình danh sách dự án / form thêm mới]"),
    ("Quản lý cư dân / người dùng\n- Bước 1: Truy cập menu Quản lý cư dân.\n- Bước 2: Xem danh sách, tìm kiếm theo tên, số điện thoại hoặc mã căn hộ.\n- Bước 3: Phê duyệt cư dân mới hoặc chỉnh sửa hồ sơ nếu cần.", "[Chèn ảnh màn hình danh sách cư dân]"),
    ("Quản lý phản ánh\n- Bước 1: Truy cập menu Quản lý phản ánh.\n- Bước 2: Xem danh sách các vấn đề được người dân gửi lên.\n- Bước 3: Click vào phản ánh để xem chi tiết, phân công người xử lý và cập nhật trạng thái (Đang xử lý, Đã hoàn thành).", "[Chèn ảnh màn hình danh sách phản ánh]"),
    ("Báo cáo và thống kê\n- Bước 1: Truy cập menu Báo cáo thống kê.\n- Bước 2: Chọn khoảng thời gian và loại báo cáo (tỷ lệ lấp đầy, tình trạng thanh toán, v.v.).\n- Bước 3: Xem biểu đồ hoặc click [Xuất Excel] để tải số liệu.", "[Chèn ảnh màn hình báo cáo]")
]

for func, demo in data1:
    row_cells = table1.add_row().cells
    row_cells[0].text = func
    row_cells[1].text = demo

set_col_widths(table1, [3.5, 3.5])
add_borders(table1)

doc.add_paragraph('\n')

# 2. Trang web dành cho người dân
p2 = doc.add_paragraph()
r2 = p2.add_run('2. Trang web dành cho người dân')
r2.bold = True
r2.font.size = Pt(12)

table2 = doc.add_table(rows=1, cols=2)
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'Chức năng/Bước thực hiện'
hdr_cells2[1].text = 'Màn hình demo'
for cell in hdr_cells2:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

data2 = [
    ("Đăng ký & Đăng nhập\n- Bước 1: Truy cập trang web hoặc tải ứng dụng.\n- Bước 2: Click [Đăng ký] (nếu chưa có tài khoản) và điền thông tin xác thực.\n- Bước 3: Nhập số điện thoại/email & mật khẩu để Đăng nhập.", "[Chèn ảnh màn hình đăng nhập / đăng ký]"),
    ("Nhận thông báo từ Ban Quản Lý\n- Bước 1: Truy cập mục Thông báo trên trang chủ.\n- Bước 2: Click vào từng thông báo để đọc nội dung chi tiết (ví dụ: cúp điện, cúp nước, nhắc nhở đóng phí).", "[Chèn ảnh màn hình danh sách thông báo]"),
    ("Tra cứu và thanh toán hóa đơn\n- Bước 1: Truy cập menu Hóa đơn của tôi.\n- Bước 2: Xem các khoản phí cần thanh toán (Phí quản lý, Gửi xe, Điện, Nước).\n- Bước 3: Chọn hóa đơn và click [Thanh toán], sau đó làm theo hướng dẫn để quét mã QR hoặc nhập thẻ.", "[Chèn ảnh màn hình tra cứu hóa đơn & thanh toán]"),
    ("Gửi phản ánh, yêu cầu hỗ trợ\n- Bước 1: Truy cập chức năng Gửi phản ánh.\n- Bước 2: Chọn loại vấn đề (An ninh, Vệ sinh, Kỹ thuật).\n- Bước 3: Điền nội dung mô tả, đính kèm hình ảnh thực tế và click [Gửi].", "[Chèn ảnh màn hình form gửi phản ánh]"),
    ("Đăng ký tiện ích nội khu\n- Bước 1: Truy cập mục Tiện ích (BBQ, Hồ bơi, Sân Tennis).\n- Bước 2: Xem lịch trống, chọn ngày giờ và số lượng người đi kèm.\n- Bước 3: Click [Đăng ký] và đợi Ban Quản Lý phê duyệt.", "[Chèn ảnh màn hình đặt lịch tiện ích]")
]

for func, demo in data2:
    row_cells = table2.add_row().cells
    row_cells[0].text = func
    row_cells[1].text = demo

set_col_widths(table2, [3.5, 3.5])
add_borders(table2)

output_path = r'c:\Users\KMSoft\OneDrive\Desktop\VGC-Smart-Draw-main\VGC-Smart-Draw-main\Huong_Dan_Su_Dung.docx'
doc.save(output_path)
print(f"File created successfully at: {output_path}")

