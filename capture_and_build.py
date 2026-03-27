import asyncio
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

async def capture_screenshots():
    screenshots = {}
    async with async_playwright() as p:
        # Tăng kích thước viewport lớn hơn để chụp full màn hình rõ nét
        browser = await p.chromium.launch(headless=True)
        page = await browser.new_page(viewport={"width": 1920, "height": 1080})
        
        # --- PHASE 1: ADMIN FLOW ---
        # 1. Màn hình đăng nhập Admin
        await page.goto("http://localhost:3000/")
        await page.wait_for_load_state('networkidle')
        await page.screenshot(path="01_admin_login.png", full_page=True)
        screenshots['admin_login'] = "01_admin_login.png"
        
        # 2. Đăng nhập thành công -> Tổng quan dự án
        await page.fill('input[placeholder="admin"]', 'admin')
        await page.fill('input[placeholder="••••••"]', '123456')
        await page.click('button:has-text("Đăng Nhập")')
        await page.wait_for_timeout(1500)
        await page.screenshot(path="02_admin_dashboard.png", full_page=True)
        screenshots['admin_dashboard'] = "02_admin_dashboard.png"

        # 3. Quản lý Quỹ Căn Hộ
        await page.click('button:has-text("Quỹ Căn Hộ")')
        await page.wait_for_timeout(1000)
        await page.screenshot(path="03_admin_inventory.png", full_page=True)
        screenshots['admin_inventory'] = "03_admin_inventory.png"
        
        # 4. Quản lý Quỹ Hồ Sơ
        await page.click('button:has-text("Quỹ Hồ Sơ")')
        await page.wait_for_timeout(1000)
        await page.screenshot(path="04_admin_profiles.png", full_page=True)
        screenshots['admin_profiles'] = "04_admin_profiles.png"

        # 5. Thiết lập Vòng Quay
        await page.click('button:has-text("Thiết lập Vòng Quay")')
        await page.wait_for_timeout(1000)
        await page.screenshot(path="05_admin_round_setup.png", full_page=True)
        screenshots['admin_round_setup'] = "05_admin_round_setup.png"

        # 6. Monitor Sự Kiện (Điều khiển)
        await page.click('button:has-text("Monitor Sự Kiện")')
        await page.wait_for_timeout(1000)
        await page.screenshot(path="06_admin_monitor.png", full_page=True)
        screenshots['admin_monitor'] = "06_admin_monitor.png"
        
        # 7. Báo cáo & Tra Cứu
        await page.click('button:has-text("Báo Cáo & Tra Cứu")')
        await page.wait_for_timeout(1000)
        await page.screenshot(path="07_admin_reports.png", full_page=True)
        screenshots['admin_reports'] = "07_admin_reports.png"


        # --- PHASE 2: USER FLOW (Mobile Viewport) ---
        user_page = await browser.new_page(viewport={"width": 600, "height": 1200}) 
        print("Navigating to User View...")
        await user_page.goto("http://localhost:3000/?role=user")
        await user_page.wait_for_load_state('networkidle')
        await user_page.wait_for_timeout(2000)
        
        # 8. User Login Screen
        print("Capturing User Login...")
        await user_page.screenshot(path="08_user_login.png", full_page=True)
        screenshots['user_login'] = "08_user_login.png"

        # 9. User Login Filled
        print("Filling User Login info...")
        # Sử dụng KC0001 thay vì KC000001
        await user_page.locator('input').nth(0).fill('KC0001')
        await user_page.locator('input[type="tel"]').fill('09123456789')
        await user_page.wait_for_timeout(500)
        await user_page.screenshot(path="09_user_login_filled.png", full_page=True)
        screenshots['user_login_filled'] = "09_user_login_filled.png"
        
        # 10. User OTP Screen
        print("Clicking NEXT...")
        await user_page.click('button:has-text("TIẾP TỤC")')
        await user_page.wait_for_timeout(2500)
        await user_page.screenshot(path="10_user_otp.png", full_page=True)
        screenshots['user_otp'] = "10_user_otp.png"
        
        # 11. User Dashboard
        print("Entering OTP...")
        # Đợi input OTP xuất hiện
        await user_page.wait_for_selector('input[placeholder="0000"]', timeout=15000)
        await user_page.fill('input[placeholder="0000"]', '1234')
        await user_page.wait_for_timeout(1000)
        
        print("Confirming OTP...")
        # Đợi nút xác thực sẵn sàng
        await user_page.click('button:has-text("XÁC")') 
        await user_page.wait_for_timeout(5000) # Đợi chuyển màn hoàn toàn
        
        # Click "VỀ TRANG CHỦ" nếu có màn thành công
        try:
            success_btn = user_page.locator('button:has-text("TRANG CHỦ")')
            if await success_btn.count() > 0:
                await success_btn.click()
                await user_page.wait_for_timeout(2000)
        except:
            pass

        print("Capturing User Dashboard...")
        await user_page.screenshot(path="11_user_dashboard.png", full_page=True)
        screenshots['user_dashboard'] = "11_user_dashboard.png"

        await browser.close()
    return screenshots

def create_docx(screenshots):
    doc = Document()

    # Cấu hình lề trang để có thêm không gian cho ảnh bự
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    title = doc.add_heading('TÀI LIỆU HƯỚNG DẪN SỬ DỤNG - VGC SMART DRAW', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_step(doc, title_text, description, img_key):
        # Add Title
        p_title = doc.add_paragraph()
        r_title = p_title.add_run(title_text)
        r_title.bold = True
        r_title.font.size = Pt(14)
        r_title.font.color.rgb = RGBColor(0, 70, 142) # #00468E
        
        # Add Description
        p_desc = doc.add_paragraph()
        r_desc = p_desc.add_run(description)
        r_desc.font.size = Pt(11)

        # Add Image
        if img_key in screenshots and os.path.exists(screenshots[img_key]):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p_img.add_run()
            # Set width to almost full page width (7.0 inches) 
            run.add_picture(screenshots[img_key], width=Inches(7.0))
        
        doc.add_page_break()

    # ----- PHẦN 1: BAN QUẢN TRỊ -----
    doc.add_heading('PHẦN 1: HƯỚNG DẪN DÀNH CHO BAN QUẢN TRỊ / CHỦ ĐẦU TƯ', level=1)
    
    add_step(doc, 
             "1.1. Đăng nhập hệ thống Quản trị (Admin)", 
             "Truy cập vào đường dẫn trang web dành cho quản trị viên.\n- Bước 1: Nhập Tên đăng nhập (Username), ví dụ: admin.\n- Bước 2: Nhập Mật khẩu.\n- Bước 3: Nhấn nút [Đăng Nhập] để vào màn hình chính.",
             "admin_login")

    add_step(doc, 
             "1.2. Màn hình Tổng quan dự án (Dashboard)", 
             "Sau khi đăng nhập thành công, hệ thống hiển thị tổng quan dự án.\n- Màn hình hiển thị tổng quỹ căn, tổng số hồ sơ tham gia.\n- Cung cấp cái nhìn nhanh về tiến độ Check-in của người dân và tỷ lệ trúng/trượt.\n- Quản trị viên có thể theo dõi biểu đồ thống kê trực quan.",
             "admin_dashboard")

    add_step(doc, 
             "1.3. Quản lý Quỹ Căn Hộ (Giỏ hàng)", 
             "Mục đích: Xem và quản lý danh sách các căn hộ được đưa ra bốc thăm.\n- Bước 1: Chọn menu [Quỹ Căn Hộ] bên thanh công cụ bên trái.\n- Bước 2: Bảng hiển thị thông tin chi tiết từng căn (Mã căn, Diện tích, Đơn giá, Trạng thái).\n- Quản trị viên có thể thêm mới, sửa đổi thông tin căn hộ hoặc Import bằng file Excel.",
             "admin_inventory")

    add_step(doc, 
             "1.4. Quản lý Quỹ Hồ Sơ (Cư dân tham gia)", 
             "Mục đích: Quản lý danh sách người dân tham gia bốc thăm.\n- Bước 1: Chọn menu [Quỹ Hồ Sơ].\n- Bước 2: Sử dụng các bộ lọc để tìm kiếm theo Mã HS, Tên, Số điện thoại hoặc Nhóm ưu tiên.\n- Có thể kiểm tra trạng thái hồ sơ: Đã duyệt, Đã nộp tiền, Đủ điều kiện.",
             "admin_profiles")

    add_step(doc, 
             "1.5. Thiết lập Vòng Quay", 
             "Mục đích: Định nghĩa các phiên bốc thăm theo từng nhóm đối tượng/tòa nhà.\n- Bước 1: Chọn menu [Thiết lập Vòng Quay].\n- Bước 2: Tạo các vòng quay mới, ví dụ: Vòng 1 cho đối tượng Ưu tiên UT1.\n- Gán danh sách căn hộ tương ứng và danh sách người tham gia vào từng vòng.",
             "admin_round_setup")

    add_step(doc, 
             "1.6. Monitor Sự Kiện (Điều hành phiên bốc thăm)", 
             "Đây là màn hình quan trọng nhất trong ngày tổ chức sự kiện.\n- Bước 1: Chọn [Monitor Sự Kiện].\n- Bước 2: Quản trị viên bật tính năng [MỞ CỔNG CHECK-IN] để cho phép người dân đăng nhập vào hệ thống.\n- Bước 3: Khi tất cả đã sẵn sàng, nhấn [KÍCH HOẠT PHIÊN BỐC THĂM] để mọi người bắt đầu quay số trên thiết bị của họ.",
             "admin_monitor")
             
    add_step(doc, 
             "1.7. Báo Cáo & Tra Cứu Kết Quả", 
             "Dùng để xem lại toàn bộ kết quả sau khi sự kiện kết thúc.\n- Bước 1: Chọn menu [Báo Cáo & Tra Cứu].\n- Bước 2: Bảng thống kê chi tiết ai đã trúng căn nào, thời gian bốc thăm.\n- Bước 3: Nhấn nút [Xuất Excel] để lưu trữ báo cáo.",
             "admin_reports")

    # ----- PHẦN 2: NGƯỜI DÂN -----
    doc.add_heading('PHẦN 2: HƯỚNG DẪN DÀNH CHO CƯ DÂN (NGƯỜI THAM GIA)', level=1)

    add_step(doc, 
             "2.1. Truy cập cổng Check-in Cư dân", 
             "Người dân sử dụng điện thoại thông minh (smartphone) truy cập vào đường link do Chủ đầu tư cung cấp.\n- Giao diện được thiết kế tối ưu cho màn hình dọc của điện thoại di động.",
             "user_login")

    add_step(doc, 
             "2.2. Nhập thông tin xác thực", 
             "Để đảm bảo tính minh bạch, người dân cần cung cấp thông tin đã đăng ký.\n- Bước 1: Cung cấp Mã Hồ Sơ của mình.\n- Bước 2: Nhập số điện thoại chính chủ.\n- Bước 3: Nhấn nút [TIẾP TỤC].",
             "user_login_filled")

    add_step(doc, 
             "2.3. Xác thực OTP", 
             "Hệ thống sẽ gửi một mã OTP gồm 4 chữ số về số điện thoại gốc.\n- Bước 1: Nhập mã OTP gồm 4 chữ số (trong bản demo nhập: 1234).\n- Bước 2: Nhấn nút xác nhận để tiến hành đăng nhập vào phòng chờ.",
             "user_otp")

    add_step(doc, 
             "2.4. Màn hình Thông tin cá nhân & Chờ bốc thăm", 
             "Sau khi đăng nhập thành công, người dân đã hoàn tất Check-in.\n- Màn hình hiển thị đầy đủ thông tin: Mã hồ sơ, Tên, Căn cước công dân.\n- Lúc này, người dân vui lòng đợi Ban tổ chức/Quản trị viên thông báo [KÍCH HOẠT PHIÊN BỐC THĂM]. Nút chức năng quay số sẽ hiện lên khi sự kiện bắt đầu.",
             "user_dashboard")

    output_path = r'c:\Users\KMSoft\OneDrive\Desktop\VGC-Smart-Draw-main\VGC-Smart-Draw-main\Huong_Dan_Su_Dung_VGC_Chi_Tiet.docx'
    doc.save(output_path)
    print(f"File created successfully at: {output_path}")

async def main():
    print("Capturing full-page screenshots for detailed manual...")
    screenshots = await capture_screenshots()
    print("Generating detailed DOCX...")
    create_docx(screenshots)
    print("Done!")

if __name__ == "__main__":
    asyncio.run(main())
