import sys
import subprocess
import os

def install_and_import(package, import_name):
    try:
        __import__(import_name)
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])
    finally:
        globals()[import_name] = __import__(import_name)

install_and_import('python-docx', 'docx')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Add Title
title = doc.add_heading('BẢNG BÁO GIÁ DỰ TOÁN - TÍNH NĂNG GIAI ĐOẠN 2 (PHASE 2)', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Dự án: VGC Smart Draw (Hệ thống Bốc thăm Nhà ở Xã hội)')
doc.add_paragraph('Phiên bản: Phase 2 (Nâng cấp Cổng thông tin & Nộp hồ sơ trực tuyến)')
doc.add_paragraph('Ngày lập báo giá: ........................')

doc.add_heading('I. TỔNG QUAN PHẠM VI CÔNG VIỆC (PHASE 2)', level=2)
doc.add_paragraph('Trong Giai đoạn 2, hệ thống tập trung mở rộng Cổng thông tin dành cho Khách hàng (User Portal) và số hóa toàn bộ quy trình tiếp nhận hồ sơ xét duyệt. Khách hàng có thể tự do tạo tài khoản, nộp hồ sơ trực tuyến, upload giấy tờ và theo dõi quá trình xét duyệt theo chuẩn ISO, thay vì chỉ đăng nhập bằng OTP để xem kết quả bốc thăm như ở Giai đoạn 1.')

doc.add_heading('II. CHI TIẾT CÁC TÍNH NĂNG NÂNG CẤP VÀ XÂY MỚI', level=2)

table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
headers = ['STT', 'Phân hệ / Tính năng', 'Mô tả chi tiết', 'Thời gian dự kiến (Man-days)', 'Đơn giá (VNĐ)', 'Thành tiền (VNĐ)']
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

data = [
    # 1. User Dashboard
    ("1", "Cổng thông tin Cá nhân (User Dashboard)", "", "", "", ""),
    ("1.1", "Đăng nhập/Đăng ký tài khoản cá nhân", "Nâng cấp hệ thống định danh. Khách hàng có thể thiết lập mật khẩu và Email.", "", "", ""),
    ("1.2", "Giao diện Bảng điều khiển", "Màn hình Tổng quan dành riêng cho từng khách hàng, hiển thị thông báo, lịch sự kiện.", "", "", ""),
    # 2. Submission Wizard
    ("2", "Quản lý Nộp Hồ sơ Trực tuyến", "", "", "", ""),
    ("2.1", "Trình hướng dẫn nộp hồ sơ (Wizard)", "Luồng nhập liệu từng bước, hướng dẫn khách kê khai mua/thuê NOXH.", "", "", ""),
    ("2.2", "Quản lý Phân loại Đối tượng", "Bổ sung logic phân loại nhóm đối tượng chính sách (K1, K2...).", "", "", ""),
    ("2.3", "Upload Giấy tờ/Tài liệu", "Hệ thống tải lên các file minh chứng (CCCD, Hộ khẩu...).", "", "", ""),
    # 3. Status Tracking
    ("3", "Tra cứu & Theo dõi Hồ sơ", "", "", "", ""),
    ("3.1", "Xem tiến độ hồ sơ", "Giao diện theo dõi trạng thái hiện tại (Đã nhận -> Đang xử lý -> Hoàn thành).", "", "", ""),
    ("3.2", "Quản lý mốc thời gian ISO", "Tính toán thời hạn nộp bản cứng và hạn chót xử lý.", "", "", ""),
    ("3.3", "Trả hồ sơ & Yêu cầu bổ sung", "Cơ chế thông báo lý do trả hồ sơ để khách cập nhật.", "", "", ""),
    # 4. Admin Modifications
    ("4", "Nâng cấp Hệ thống Quản trị", "", "", "", ""),
    ("4.1", "Log lịch sử xử lý hồ sơ", "Bảng theo dõi các thao tác trên hồ sơ của khách (Audit Trail).", "", "", ""),
    ("4.2", "Cấu hình thời gian tự động", "Bổ sung thiết lập thời gian cho GSV và tự động chuyển vòng.", "", "", ""),
]

for item in data:
    row_cells = table.add_row().cells
    for i in range(6):
        row_cells[i].text = item[i]
        if "." not in item[0] and item[0] != "":
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

for row in table.rows:
    row.cells[0].width = Inches(0.5)
    row.cells[1].width = Inches(1.5)
    row.cells[2].width = Inches(2.5)
    row.cells[3].width = Inches(0.8)
    row.cells[4].width = Inches(1.0)
    row.cells[5].width = Inches(1.0)

doc.add_heading('III. TỔNG CỘNG CHI PHÍ', level=2)
doc.add_paragraph('Tổng chi phí dự kiến: ............................ VNĐ')
doc.add_paragraph('Thuế VAT (10%): ............................ VNĐ')
p = doc.add_paragraph()
r = p.add_run('TỔNG GIÁ TRỊ THANH TOÁN: ............................ VNĐ')
r.bold = True

doc.add_paragraph('\nGhi chú:\n- Bảng báo giá trên chưa bao gồm chi phí nâng cấp Server / Lưu trữ File.\n- Thời gian triển khai tính từ ngày hai bên ký kết hợp đồng.')

output_path = r'c:\Users\KMSoft\OneDrive\Desktop\Phase-2-NOXH-Smart-Draw--HAN-main\Bao_Gia_Phase_2.docx'
doc.save(output_path)
print(f"Bao gia saved to {output_path}")
