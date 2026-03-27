import asyncio
import os
import sys
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# Configuration
BASE_URL = "http://localhost:3003"
OUTPUT_DOCX = "Huong_Dan_Su_Dung_VGC_Smart_Draw_Pro.docx"
IMAGE_DIR = "captured_images_pro"

# Ensure UTF-8 for printing if possible, but we'll stay with English logs
if not os.path.exists(IMAGE_DIR):
    os.makedirs(IMAGE_DIR)

async def capture_screen(page, name, wait_selector=None):
    if wait_selector:
        try:
            await page.wait_for_selector(wait_selector, timeout=5000)
        except:
            print(f"Warning: Timeout waiting for selector {wait_selector}")
    
    await asyncio.sleep(1) # Settle animations
    path = os.path.join(IMAGE_DIR, f"{name}.png")
    await page.screenshot(path=path, full_page=True)
    return path

async def close_modal(page):
    try:
        # Try clicking anywhere outside or use Escape
        await page.keyboard.press("Escape")
        await asyncio.sleep(0.5)
        # Also try clicking a generic close button if visible
        for btn in await page.locator("button").all():
            text = await btn.inner_text()
            if not text and await btn.locator("svg.lucide-x").is_visible():
                await btn.click(force=True)
                break
    except:
        pass
    await asyncio.sleep(1)

def add_step_to_docx(doc, title, description, image_path):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.3)
    
    cells = table.rows[0].cells
    cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    p_title = cells[0].paragraphs[0]
    run_title = p_title.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(12)
    
    desc_p = cells[0].add_paragraph(description)
    desc_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    p_img = cells[1].paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(image_path, width=Inches(4.1))

    doc.add_paragraph()

async def run():
    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=False)
        context = await browser.new_context(viewport={'width': 1920, 'height': 1080})
        page = await context.new_page()

        steps = []

        # --- ADMIN FLOW ---
        print("Starting Admin Flow...")
        await page.goto(f"{BASE_URL}?role=admin")
        await page.wait_for_selector("input", timeout=10000)
        
        # Capture Login Screen
        path_login = await capture_screen(page, "0_Admin_Login")
        steps.append(("0. Trang đăng nhập Quản trị", "Giao diện đăng nhập bảo mật dành cho Quản trị viên hệ thống VGC Smart Draw.", path_login))

        await page.locator("input").nth(0).fill("admin")
        await page.locator("input").nth(1).fill("123456")
        await page.click("button:has-text('NHẬP')") # Simplified selector
        
        await asyncio.sleep(2)
        path = await capture_screen(page, "1_Admin_Dashboard")
        steps.append(("1. Bàn điều khiển Admin", "Giao diện chính dành cho quản trị viên, hiển thị tổng quan các thông số dự án, thống kê hồ sơ và trạng thái hệ thống.", path))

        # Navigating through sidebar
        tabs = [
            ("Quỹ Hồ Sơ", "2_Admin_Profiles", "2. Quản lý Hồ sơ", "Danh sách toàn bộ hồ sơ khách hàng tham gia bốc thăm kèm theo các công cụ lọc và tìm kiếm thông minh."),
            ("Quỹ Căn Hộ", "4_Admin_Inventory", "4. Quản lý Quỹ căn", "Theo dõi tình trạng căn hộ (Chưa chủ / Có chủ) và cấu hình quyền (Mua/Thuê/Thuê-Mua) cho từng căn."),
            ("Thiết lập Vòng Quay", "6_Admin_Rounds", "6. Cấu hình Vòng quay", "Thiết lập các tham số cho từng đợt bốc thăm: Thời gian bắt đầu/kết thúc, nhóm đối tượng và số lượng căn hộ áp dụng."),
            ("Monitor Sự Kiện", "8_Admin_Monitor", "8. Giám sát thời gian thực", "Màn hình điều hành trung tâm cho phép theo dõi số người đang online, trạng thái kết nối và tiến độ bốc thăm."),
            ("Báo Cáo & Tra Cứu", "9_Admin_Reports", "9. Báo cáo kết quả bốc thăm", "Hệ thống tự động tổng hợp kết quả, hỗ trợ tra cứu nhanh và trích xuất dữ liệu ra file Excel phục vụ lưu trữ."),
            ("Tài Khoản Hệ Thống", "10_Admin_Accounts", "10. Tài khoản & Phân quyền", "Quản lý danh sách tài khoản quản trị, quy định vai trò (Roles) và tra cứu nhật ký hệ thống.")
        ]

        for tab_label, img_name, step_title, step_desc in tabs:
            print(f"Navigating to tab...") # Removed non-English content
            await page.click(f"button:has-text('{tab_label}')", force=True)
            path = await capture_screen(page, img_name)
            steps.append((step_title, step_desc, path))
            
            if tab_label == "Tài Khoản Hệ Thống":
                # Capture sub-tabs: Tài khoản, Phân quyền, Nhật ký
                sub_tabs = [
                    ("Phân quyền", "8b_Admin_Permissions", "8b. Quản lý Vai trò & Phân quyền", "Thiết lập quyền hạn chi tiết cho từng nhóm vai trò trên từng tính năng của hệ thống."),
                    ("Nhật ký", "8c_Admin_Logs", "8c. Nhật ký hoạt động", "Lưu trữ lịch sử thao tác của các tài khoản quản trị để phục vụ công tác đối soát và kiểm tra.")
                ]
                for sub_label, sub_img, sub_title, sub_desc in sub_tabs:
                    print(f"Navigating to sub-tab...") # Removed non-English content
                    try:
                        await page.click(f"button:has-text('{sub_label}')", force=True)
                        path_sub = await capture_screen(page, sub_img)
                        steps.append((sub_title, sub_desc, path_sub))
                    except: pass
            
            # Special actions for some tabs
            if tab_label in ["Quỹ Hồ Sơ", "Quỹ Căn Hộ"]:
                try:
                    await page.click("button:has-text('Import')", force=True)
                    path_mod = await capture_screen(page, f"mod_{img_name}")
                    steps.append((f"{step_title} - Nạp dữ liệu", f"Tính năng nạp dữ liệu hàng loạt từ file Excel cho {tab_label}.", path_mod))
                    await close_modal(page)
                except: pass
            if tab_label == "Thiết lập Vòng Quay":
                try:
                    await page.click("button:has-text('Thêm')", force=True)
                    path_mod = await capture_screen(page, "7_Admin_Add_Round")
                    steps.append(("7. Thêm mới Vòng quay", "Giao diện khởi tạo vòng quay mới với các tùy chọn lọc đối tượng và căn hộ linh hoạt.", path_mod))
                    await close_modal(page)
                except: pass

        # --- USER FLOW ---
        print("Starting User Flow...")
        await page.goto(f"{BASE_URL}?role=user")
        await page.wait_for_timeout(2000)
        path = await capture_screen(page, "10_User_Login")
        steps.append(("10. Trang đăng nhập", "Cổng đăng nhập an toàn dành cho khách hàng tham gia bốc thăm bằng Mã hồ sơ và Số điện thoại.", path))
        
        await page.locator("input").nth(0).fill("KC0003")
        await page.locator("input").nth(1).fill("09123456787")
        await page.click("button:has-text('TỤC')") # TIẾP TỤC
        
        await page.wait_for_timeout(1000)
        path = await capture_screen(page, "11_User_OTP")
        steps.append(("11. Xác thực OTP", "Hệ thống gửi mã xác thực về điện thoại khách hàng để đảm bảo tính bảo mật và định danh chính xác.", path))
        await page.locator("input").nth(0).fill("1234")
        await page.click("button:has-text('THỰC')") # XÁC THỰC
        await page.wait_for_timeout(2000)
        
        path = await capture_screen(page, "12_User_Dashboard")
        steps.append(("12. Trang chủ Người dùng", "Sau khi đăng nhập thành công, khách hàng có thể theo dõi trạng thái cá nhân và lối vào phiên bốc thăm.", path))
        
        # Guide
        try:
            await page.click("button:has-text('Học')", force=True) # Maybe 'Hướng dẫn'? Let's use simpler
            await page.click("button:has-text('Thể lệ')", force=True)
            await asyncio.sleep(1)
            path = await capture_screen(page, "13_User_Rules")
            steps.append(("13. Quy định & Hướng dẫn", "Cung cấp đầy đủ thông tin về quy trình bốc thăm và các bước thực hiện để đảm bảo quyền lợi khách hàng.", path))
            await page.go_back()
            await page.wait_for_timeout(1000)
        except: pass

        # Activate Demo if needed
        try:
            demo_btn = page.locator("button:has-text('Demo')")
            if await demo_btn.is_visible():
                await demo_btn.click()
                await asyncio.sleep(0.5)
        except: pass

        # Enter Lobby
        print("Entering Lobby...")
        await page.click("button:has-text('NGAY')", force=True) # VÀO BỐC THĂM NGAY
        await asyncio.sleep(1)
        path = await capture_screen(page, "14_User_Lobby")
        steps.append(("14. Phòng chờ bốc thăm", "Giao diện chuẩn bị trước thời điểm chính thức bắt đầu phiên quay số.", path))

        # Start Session Demo if needed
        try:
            start_demo = page.locator("button:has-text('Bắt đầu')")
            if await start_demo.is_visible():
                await start_demo.click()
                await asyncio.sleep(0.5)
        except: pass

        # Spin
        print("Trying to spin...")
        try:
            spin_btn = page.locator("button:has-text('Quay')")
            await spin_btn.wait_for(state="visible", timeout=5000)
            path = await capture_screen(page, "15_User_Live_Ready")
            steps.append(("15. Kích hoạt lượt quay", "Khi phiên bốc thăm mở, nút 'Quay' sẽ hiện ra để khách hàng trực tiếp tham gia bốc thăm căn hộ.", path))

            await spin_btn.click()
            await asyncio.sleep(0.5)
            path = await capture_screen(page, "16_User_Spinning")
            steps.append(("16. Quá trình bốc thăm", "Hệ thống thực hiện vòng quay ngẫu nhiên được trình diễn bằng hiệu ứng hình ảnh sống động.", path))

            await page.wait_for_selector("text=CHÚC MỪNG", timeout=15000)
            path = await capture_screen(page, "17_User_Result")
            steps.append(("17. Thông báo kết quả", "Kết quả được hiển thị ngay lập tức, thông báo cụ thể về quyền sở hữu căn hộ.", path))

            await page.click("button:has-text('XEM')", force=True) # XEM CHỨNG NHẬN
            await asyncio.sleep(1)
            path = await capture_screen(page, "18_User_Certificate")
            steps.append(("18. Chứng nhận điện tử", "Khách hàng nhận được chứng nhận quyền chọn căn hộ chính thức với mã QR tra cứu duy nhất.", path))
        except Exception as e:
            print(f"Skipping spin steps: {e}")

        await browser.close()

        # Build DOCX
        print("Building final document...")
        doc = Document()
        doc.add_heading('HƯỚNG DẪN SỬ DỤNG HỆ THỐNG VGC SMART DRAW', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Tài liệu hướng dẫn chi tiết dành cho Quản trị viên và Người dùng").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        for title, desc, img in steps:
            add_step_to_docx(doc, title, desc, img)

        doc.save(OUTPUT_DOCX)
        print(f"Manual successfully generated: {OUTPUT_DOCX}")

if __name__ == "__main__":
    asyncio.run(run())
